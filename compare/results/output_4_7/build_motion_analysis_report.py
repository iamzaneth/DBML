from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
OUT = BASE / "bao_cao_phan_tich_motion_4_7_ASL_VSL.docx"


COLORS = {
    "blue": RGBColor(46, 116, 181),
    "dark_blue": RGBColor(31, 77, 120),
    "ink": RGBColor(0, 0, 0),
    "muted": RGBColor(90, 90, 90),
    "header_fill": "F2F4F7",
    "callout_fill": "F4F6F9",
}


def fmt(x, digits=2):
    if pd.isna(x):
        return "NA"
    if abs(float(x)) >= 100:
        return f"{float(x):,.1f}"
    return f"{float(x):,.{digits}f}"


def pct(x, digits=1):
    if pd.isna(x):
        return "NA"
    return f"{float(x) * 100:.{digits}f}%"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color=COLORS["ink"], size=font_size)
        set_cell_shading(hdr[i], COLORS["header_fill"])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=font_size)
    set_table_widths(table, widths)
    doc.add_paragraph()
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, COLORS["callout_fill"])
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = COLORS["dark_blue"]
    r.font.size = Pt(10.5)
    p.add_run("\n" + body).font.size = Pt(10)
    doc.add_paragraph()


def add_picture(doc, filename, caption, width=6.25):
    path = BASE / "figures" / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = COLORS["muted"]


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, COLORS["blue"], 16, 8),
        ("Heading 2", 13, COLORS["blue"], 12, 6),
        ("Heading 3", 12, COLORS["dark_blue"], 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    for name in ["List Bullet", "List Number"]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.paragraph_format.left_indent = Inches(0.5)
        st.paragraph_format.first_line_indent = Inches(-0.25)
        st.paragraph_format.space_after = Pt(8)
        st.paragraph_format.line_spacing = 1.167


def metric_value(summary, metric, dataset, field="mean"):
    row = summary[(summary["metric"] == metric) & (summary["dataset"] == dataset)]
    return float(row.iloc[0][field])


def main():
    summary = pd.read_csv(BASE / "motion_summary_by_dataset.csv")
    wide = pd.read_csv(BASE / "motion_summary_by_dataset_wide.csv").set_index("metric")
    labels = pd.read_csv(BASE / "motion_features_by_label.csv")
    top = pd.read_csv(BASE / "top_complex_motion_labels.csv")
    seq = pd.read_csv(BASE / "sequence_length_by_label.csv")
    dtw_sum = pd.read_csv(BASE / "dtw" / "dtw_intra_class_summary.csv")
    dtw_pairs = pd.read_csv(BASE / "dtw" / "dtw_inter_class_candidate_pairs.csv")

    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("Báo cáo phân tích Motion 4.7 cho ASL và VSL")
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = COLORS["dark_blue"]

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    sr = subtitle.add_run(
        "Nguồn: C:\\Windows_Application\\DBML\\compare\\results\\output_4_7. "
        "Mục tiêu: hiểu đặc điểm chuyển động và tác động khi train mô hình nhận dạng ký hiệu."
    )
    sr.font.size = Pt(10)
    sr.font.color.rgb = COLORS["muted"]

    add_callout(
        doc,
        "Kết luận ngắn",
        "VSL có chuyển động dài hơn, nhanh hơn và cong/phức tạp hơn ASL; tuy nhiên VSL lại thiếu mẫu theo nhãn nghiêm trọng. "
        "Nếu train phân loại trực tiếp trên toàn bộ nhãn VSL hiện tại, mô hình rất dễ overfit, validation không ổn định và nhiều lớp không thể học được ranh giới tin cậy. "
        "ASL có nhiều mẫu hơn mỗi nhãn nhưng vẫn có nhiều nhãn ít mẫu, nhãn thiếu tay và một số cặp chuyển động rất giống nhau."
    )

    doc.add_heading("1. Phạm vi dữ liệu và chỉ số", level=1)
    add_table(
        doc,
        ["Hạng mục", "ASL", "VSL", "Ý nghĩa khi train"],
        [
            [
                "Số video/mẫu",
                f"{int(metric_value(summary, 'num_frames', 'ASL', 'count')):,}",
                f"{int(metric_value(summary, 'num_frames', 'VSL', 'count')):,}",
                "ASL có nhiều mẫu hơn; VSL ít mẫu hơn nhưng nhiều nhãn hơn, gây mất cân bằng nặng.",
            ],
            [
                "Số nhãn",
                f"{len(labels[labels.dataset == 'ASL']):,}",
                f"{len(labels[labels.dataset == 'VSL']):,}",
                "Số nhãn lớn làm bài toán khó; nhãn có quá ít mẫu không đủ để học mô hình giám sát thông thường.",
            ],
            [
                "Median mẫu/nhãn",
                fmt(labels[labels.dataset == "ASL"].num_samples.median(), 1),
                fmt(labels[labels.dataset == "VSL"].num_samples.median(), 1),
                "VSL median chỉ 1 mẫu/nhãn, không đủ chia train/validation/test theo lớp.",
            ],
            [
                "Tỷ lệ nhãn < 5 mẫu",
                f"{int((labels[labels.dataset == 'ASL'].num_samples < 5).sum()):,}",
                f"{int((labels[labels.dataset == 'VSL'].num_samples < 5).sum()):,}",
                "Các nhãn này nên gom vào nhóm few-shot/không train classifier chuẩn hoặc cần bổ sung dữ liệu.",
            ],
        ],
        [1.45, 1.0, 1.0, 3.05],
        font_size=8.5,
    )

    doc.add_heading("2. Descriptive - Điều gì đã xảy ra?", level=1)
    doc.add_paragraph(
        "Phân tích motion cho thấy VSL có đặc trưng chuyển động mạnh hơn ASL trên hầu hết chỉ số động học. "
        "Độ dài chuỗi trung bình của VSL là 92.48 frame, cao hơn ASL 65.65 frame khoảng 1.41 lần. "
        "Tổng quãng đường tay và trajectory_length của VSL đều gấp khoảng 2.07 lần ASL. "
        "Mean velocity của VSL cao hơn 1.51 lần, motion variance cao hơn 2.17 lần và số lần đổi hướng cao hơn 1.65 lần."
    )
    doc.add_paragraph(
        "Ngược lại, straightness_ratio của VSL chỉ bằng khoảng 0.28 lần ASL, nghĩa là quỹ đạo VSL ít thẳng hơn, có nhiều đường cong, lặp hoặc đổi hướng. "
        "Displacement của VSL thấp hơn ASL dù tổng chuyển động cao hơn, cho thấy nhiều ký hiệu VSL có chuyển động vòng, dao động hoặc quay về gần vị trí ban đầu thay vì đi theo đường thẳng."
    )
    add_table(
        doc,
        ["Metric", "ASL mean", "VSL mean", "VSL/ASL", "Diễn giải"],
        [
            ["num_frames", fmt(wide.loc["num_frames", "ASL"]), fmt(wide.loc["num_frames", "VSL"]), "1.41x", "VSL dài hơn, cần xử lý chuỗi thời gian tốt hơn."],
            ["total_motion", fmt(wide.loc["total_motion", "ASL"]), fmt(wide.loc["total_motion", "VSL"]), "2.07x", "VSL có biên độ/quãng đường chuyển động lớn hơn."],
            ["mean_velocity", fmt(wide.loc["mean_velocity", "ASL"]), fmt(wide.loc["mean_velocity", "VSL"]), "1.51x", "Tốc độ ký hiệu VSL cao hơn, dễ nhạy với sampling/fps."],
            ["motion_variance", fmt(wide.loc["motion_variance", "ASL"], 3), fmt(wide.loc["motion_variance", "VSL"], 3), "2.17x", "Chuyển động VSL kém ổn định hơn theo thời gian."],
            ["direction_change_count", fmt(wide.loc["direction_change_count", "ASL"]), fmt(wide.loc["direction_change_count", "VSL"]), "1.65x", "VSL có nhiều đổi hướng hơn."],
            ["straightness_ratio", fmt(wide.loc["straightness_ratio", "ASL"], 3), fmt(wide.loc["straightness_ratio", "VSL"], 3), "0.28x", "VSL có quỹ đạo cong/lặp nhiều hơn."],
            ["hand_missing_ratio", pct(wide.loc["hand_frame_missing_ratio", "ASL"]), pct(wide.loc["hand_frame_missing_ratio", "VSL"]), "0.81x", "ASL thiếu tay nhiều hơn trung bình, nhưng VSL có một số nhãn lỗi nặng."],
        ],
        [1.55, 0.9, 0.9, 0.8, 2.95],
        font_size=8.2,
    )
    add_picture(doc, "boxplot_total_motion.png", "Hình 1. Phân bố total_motion: VSL lệch lên cao hơn ASL.", width=5.9)
    add_picture(doc, "boxplot_straightness_ratio.png", "Hình 2. Straightness ratio thấp hơn ở VSL, phản ánh quỹ đạo phức tạp hơn.", width=5.9)

    doc.add_heading("3. Diagnostic - Tại sao xảy ra?", level=1)
    add_bullet(
        doc,
        "Khác biệt về động học: VSL trong bộ dữ liệu này có chuỗi dài hơn, nhiều đổi hướng hơn và khoảng cách hai tay lớn hơn. Vì vậy motion variance, trajectory length và velocity đều cao."
    )
    add_bullet(
        doc,
        "Chuyển động không tuyến tính: tổng quãng đường VSL cao nhưng displacement thấp hơn ASL. Điều này thường xảy ra khi tay đi theo vòng, lặp nhịp hoặc trở lại gần vị trí ban đầu."
    )
    add_bullet(
        doc,
        "Mất cân bằng dữ liệu: ASL có 2,000 nhãn và 11,980 mẫu, median 6 mẫu/nhãn. VSL có 3,314 nhãn nhưng chỉ 4,362 mẫu, median 1 mẫu/nhãn; gần như toàn bộ nhãn VSL có ít hơn 5 mẫu."
    )
    add_bullet(
        doc,
        "Chất lượng landmark không đồng đều: hand_frame_missing_ratio trung bình là 39.6% ở ASL và 32.3% ở VSL. Một số nhãn có tỷ lệ thiếu tay rất cao, làm đặc trưng chuyển động bị nhiễu hoặc bằng 0/NaN."
    )
    doc.add_paragraph(
        "DTW bổ sung góc nhìn về tính nhất quán theo nhãn. ASL có intra-class DTW trung bình 0.7397, VSL có 0.5087 trên các nhãn đủ mẫu để tính. "
        "Con số VSL thấp hơn không nên hiểu là VSL dễ hơn, vì chỉ 551/3,314 nhãn VSL có đủ mẫu để tính intra-class DTW; phần còn lại gần như là single-shot nên không thể đo độ biến thiên trong lớp."
    )
    add_table(
        doc,
        ["Nhóm rủi ro", "Dấu hiệu trong kết quả", "Tác động"],
        [
            ["Few-shot VSL", "3,313/3,314 nhãn VSL có < 5 mẫu; median 1 mẫu/nhãn.", "Classifier nhiều lớp sẽ học thuộc mẫu train, validation theo lớp không đáng tin."],
            ["Nhãn chuyển động phức tạp", "VSL có total_motion, velocity, variance và đổi hướng cao hơn.", "Cần mô hình temporal đủ mạnh; mô hình dùng mean feature đơn giản sẽ mất thông tin thứ tự."],
            ["Nhãn dễ nhầm", "DTW inter-class có nhiều cặp rất thấp, ASL min 0.0522; VSL có cặp 0.0000.", "Mô hình dễ nhầm các nhãn có quỹ đạo tay giống nhau nếu thiếu hand shape/pose/face context."],
            ["Thiếu landmark tay", "Một số nhãn ASL/VSL có hand_missing_mean rất cao.", "Đặc trưng đầu vào bị mất, train dễ học nhiễu hoặc tạo bias theo video/detector."],
        ],
        [1.45, 2.55, 2.5],
        font_size=8.5,
    )

    doc.add_heading("4. Predictive - Điều gì sẽ xảy ra khi train mô hình?", level=1)
    add_callout(
        doc,
        "Dự báo chính",
        "Nếu train một mô hình phân loại nhiều lớp trực tiếp trên toàn bộ nhãn hiện có, kết quả tổng thể có thể nhìn ổn trên split ngẫu nhiên theo video nhưng sẽ kém bền khi đánh giá theo nhãn/người ký mới. "
        "Rủi ro lớn nhất nằm ở VSL: quá nhiều nhãn chỉ có 1 mẫu nên không thể học biến thiên trong lớp."
    )
    add_bullet(doc, "ASL: có khả năng train baseline supervised được, nhưng accuracy sẽ bị kéo xuống bởi nhãn ít mẫu, nhãn thiếu tay và các cặp có quỹ đạo tương tự như cute/sweet, sugar/voice, u/v.")
    add_bullet(doc, "VSL: train classifier toàn bộ 3,314 nhãn hiện tại sẽ rất dễ overfit. Các nhãn 1 mẫu không thể xuất hiện đồng thời ở train và validation theo cách hợp lệ, nên macro-F1 sẽ thấp hoặc không ổn định.")
    add_bullet(doc, "Các ký hiệu chuyển động lớn/nhanh sẽ cần chuẩn hóa tốc độ và resampling; nếu không, mô hình học khác biệt độ dài video thay vì bản chất ký hiệu.")
    add_bullet(doc, "Các cặp DTW thấp sẽ là nguồn nhầm lẫn chính. Nếu chỉ dùng tọa độ tay, mô hình thiếu thông tin hand shape, vị trí so với mặt/cơ thể và ngữ cảnh hai tay để phân biệt.")
    add_picture(doc, "histogram_intra_class_dtw.png", "Hình 3. Intra-class DTW cho thấy độ biến thiên trong cùng nhãn.", width=5.9)
    add_picture(doc, "top_inter_class_dtw_pairs_VSL.png", "Hình 4. Một số cặp VSL có quỹ đạo gần như trùng nhau theo DTW.", width=6.1)

    doc.add_heading("5. Prescriptive - Nên làm gì để train tốt hơn?", level=1)
    doc.add_heading("5.1 Chuẩn bị dữ liệu", level=2)
    add_number(doc, "Không train toàn bộ VSL như bài toán 3,314 lớp ở giai đoạn đầu. Tạo tập con nhãn đủ mẫu, ví dụ chỉ giữ nhãn có >= 5 hoặc >= 10 video; phần còn lại tách thành few-shot/open-set.")
    add_number(doc, "Dùng split theo nhãn và theo người ký nếu có metadata. Tránh split ngẫu nhiên làm rò rỉ phong cách ký hoặc video gần giống nhau giữa train và validation.")
    add_number(doc, "Lọc hoặc gắn cờ mẫu có hand_missing_ratio quá cao. Ngưỡng thực dụng: loại mẫu > 0.7; kiểm tra thủ công nhóm 0.5-0.7; giữ cờ missing_ratio làm feature phụ hoặc sample weight.")
    add_number(doc, "Chuẩn hóa độ dài chuỗi: resample/pad/mask về số frame thống nhất, nhưng vẫn giữ mask thật để mô hình biết frame nào hợp lệ.")
    add_number(doc, "Chuẩn hóa tọa độ theo cơ thể: scale theo vai/torso, center theo body/pose, và dùng tọa độ tương đối hand-face, hand-body, left-right hand để giảm nhiễu do camera/người ký.")

    doc.add_heading("5.2 Thiết kế mô hình", level=2)
    add_bullet(doc, "Baseline nên có hai tầng: baseline thống kê với mean/std/velocity để kiểm tra nhanh, sau đó mô hình temporal như BiLSTM/GRU, TCN hoặc Transformer nhẹ với attention mask.")
    add_bullet(doc, "Với VSL ít mẫu, ưu tiên metric learning/prototypical network, contrastive learning hoặc pretrain trên ASL rồi fine-tune VSL thay vì train classifier lớn từ đầu.")
    add_bullet(doc, "Kết hợp nhiều nhóm đặc trưng: hand landmarks, pose/body, khoảng cách hai tay, khoảng cách tay-mặt/tay-thân, velocity/acceleration. Không chỉ dùng total_motion vì dễ nhầm các nhãn cùng quỹ đạo.")
    add_bullet(doc, "Dùng augmentation có kiểm soát: jitter nhẹ tọa độ, time-warp, random crop/pad temporal, mirror nếu hợp lý về ngôn ngữ ký hiệu. Không augmentation quá mạnh làm đổi nghĩa ký hiệu.")

    doc.add_heading("5.3 Chiến lược đánh giá", level=2)
    add_bullet(doc, "Báo cáo macro-F1, balanced accuracy và top-k accuracy, không chỉ accuracy tổng vì dữ liệu lệch nhãn rất mạnh.")
    add_bullet(doc, "Tạo confusion report riêng cho các cặp inter-class DTW thấp để biết mô hình có thật sự phân biệt được nhãn giống quỹ đạo không.")
    add_bullet(doc, "Theo dõi metric theo nhóm: nhãn ít mẫu, nhãn motion_complexity_score cao, nhãn hand_missing cao, nhãn sequence_length_variance cao.")
    add_bullet(doc, "So sánh ít nhất ba chế độ: ASL-only, VSL đủ mẫu, và transfer ASL -> VSL. Không trộn ASL/VSL một cách mù nếu mục tiêu cuối là nhận dạng VSL.")

    doc.add_heading("6. Danh sách ưu tiên hành động", level=1)
    add_table(
        doc,
        ["Ưu tiên", "Việc cần làm", "Lý do", "Kết quả mong đợi"],
        [
            ["1", "Tạo subset train được: nhãn >= 5 hoặc >= 10 mẫu.", "VSL median 1 mẫu/nhãn, không đủ supervised split.", "Validation hợp lệ hơn, macro-F1 có ý nghĩa hơn."],
            ["2", "Lọc/gắn cờ mẫu thiếu tay cao.", "Missing landmark làm motion feature sai hoặc NaN.", "Giảm nhiễu đầu vào và giảm nhầm lẫn do detector."],
            ["3", "Resample + mask chuỗi thời gian.", "VSL dài hơn và variance frame lớn hơn.", "Mô hình học động tác thay vì học độ dài video."],
            ["4", "Train baseline temporal + metric learning.", "Motion phức tạp và dữ liệu VSL few-shot.", "Tăng khả năng tổng quát cho nhãn ít mẫu."],
            ["5", "Đánh giá theo nhóm khó và cặp DTW thấp.", "Nhiều cặp có quỹ đạo tương tự.", "Biết rõ lỗi nằm ở dữ liệu, feature hay kiến trúc mô hình."],
        ],
        [0.65, 2.0, 2.0, 1.85],
        font_size=8.2,
    )

    doc.add_heading("7. Nhãn/cặp cần theo dõi", level=1)
    asl_complex = top[top.dataset == "ASL"].head(8)
    vsl_complex = top[top.dataset == "VSL"].head(8)
    rows = []
    for _, row in asl_complex.iterrows():
        rows.append(["ASL", row["label"], int(row["num_samples"]), fmt(row["mean_total_motion"]), fmt(row["motion_complexity_score"], 3)])
    for _, row in vsl_complex.iterrows():
        rows.append(["VSL", row["label"], int(row["num_samples"]), fmt(row["mean_total_motion"]), fmt(row["motion_complexity_score"], 3)])
    add_table(doc, ["Dataset", "Nhãn phức tạp", "Mẫu", "Total motion", "Complexity"], rows, [0.8, 2.3, 0.65, 1.15, 1.1], font_size=8.5)

    pair_rows = []
    for ds in ["ASL", "VSL"]:
        sample = dtw_pairs[dtw_pairs.dataset == ds].sort_values("dtw_distance").head(6)
        for _, row in sample.iterrows():
            pair_rows.append([ds, row["label_a"], row["label_b"], fmt(row["dtw_distance"], 4)])
    add_table(doc, ["Dataset", "Label A", "Label B", "DTW distance"], pair_rows, [0.75, 2.3, 2.3, 1.15], font_size=8.2)

    doc.add_heading("8. Kết luận", level=1)
    doc.add_paragraph(
        "Kết quả motion 4.7 không chỉ mô tả sự khác biệt ASL/VSL mà còn chỉ ra rủi ro train rất cụ thể. "
        "VSL là bài toán khó hơn về động học vì chuyển động dài, nhanh, cong và nhiều đổi hướng; đồng thời khó hơn về dữ liệu vì số mẫu trên mỗi nhãn quá thấp. "
        "Do đó giải pháp tốt nhất không phải là train ngay một classifier lớn trên toàn bộ nhãn, mà là làm sạch dữ liệu, chọn subset đủ mẫu, chuẩn hóa chuỗi, dùng mô hình temporal có mask và áp dụng transfer/metric learning cho VSL."
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Báo cáo tạo từ kết quả motion_analysis_4_7_colab - output_4_7")
    fr.font.size = Pt(8)
    fr.font.color.rgb = COLORS["muted"]

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
from pathlib import Path
import math

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
OUT = BASE / "bao_cao_phan_tich_motion_4_7_ASL_VSL.docx"
FIG = BASE / "figures"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(85, 85, 85)
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"


def read_csv(rel):
    return pd.read_csv(BASE / rel, encoding="utf-8-sig")


def fmt(v, digits=3):
    if v is None:
        return ""
    try:
        if pd.isna(v):
            return ""
    except TypeError:
        pass
    if isinstance(v, str):
        return v
    if isinstance(v, (int,)) or (isinstance(v, float) and float(v).is_integer() and abs(v) >= 100):
        return f"{int(v):,}"
    return f"{float(v):,.{digits}f}"


def pct(v):
    if pd.isna(v):
        return ""
    return f"{float(v) * 100:.1f}%"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_in):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
                set_cell_margins(row.cells[idx])
                row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None, header_fill=LIGHT_GRAY, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        p = hdr[i].paragraphs[0]
        r = p.add_run(str(h))
        r.bold = True
        r.font.size = Pt(font_size)
        set_cell_shading(hdr[i], header_fill)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(str(value))
            r.font.size = Pt(font_size)
    if widths:
        set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_para(doc, text="", bold=False, italic=False, color=None, size=11, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.167
        p.add_run(item)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = BLUE if level <= 2 else DARK_BLUE
        r.font.name = "Calibri"
    return p


def add_figure(doc, filename, caption, width=6.15):
    path = FIG / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("title", filename)
    doc_pr.set("descr", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "Section 4.7 Motion Analysis - ASL vs VSL"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = GRAY
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "Generated from compare/results/output_4_7"
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = GRAY
    return doc


def main():
    video = read_csv("motion_features_per_video.csv")
    by_label = read_csv("motion_features_by_label.csv")
    seq = read_csv("sequence_length_by_label.csv")
    top_complex = read_csv("top_complex_motion_labels.csv")
    intra = read_csv("dtw/dtw_intra_class_summary.csv")
    inter = read_csv("dtw/dtw_inter_class_candidate_pairs.csv")
    wide = read_csv("motion_summary_by_dataset_wide.csv")
    debug = read_csv("debug/npz_structure_inspection.csv")

    doc = setup_doc()

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("Báo cáo phân tích chuyển động 4.7: ASL và VSL")
    title_run.bold = True
    title_run.font.size = Pt(23)
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(14)
    r = sub.add_run("Tổng hợp CSV, biểu đồ và khuyến nghị huấn luyện mô hình học máy từ output_4_7")
    r.font.size = Pt(13)
    r.font.color.rgb = GRAY

    meta_rows = [
        ("Nguồn kết quả", str(BASE)),
        ("Quy mô dữ liệu", f"{len(video):,} video; {by_label['label'].nunique():,} nhãn duy nhất theo tên; {len(by_label):,} dòng nhãn-theo-dataset"),
        ("CSV chính", "motion_features_per_video, motion_features_by_label, sequence_length_by_label, motion_summary_by_dataset, DTW intra/inter"),
        ("Biểu đồ", f"{len(list(FIG.glob('*.png')))} PNG trong thư mục figures"),
    ]
    add_table(doc, ["Mục", "Giá trị"], meta_rows, widths=[1.55, 4.95], header_fill=LIGHT_BLUE, font_size=9)

    add_heading(doc, "1. Kết luận nhanh", 1)
    add_para(
        doc,
        "Kết quả cho thấy VSL không chỉ khác ASL về ngôn ngữ ký hiệu mà còn khác mạnh về phân phối chuyển động, độ dài chuỗi và mật độ mẫu theo nhãn. Vì vậy không nên train chung ASL/VSL như một tập đồng nhất nếu không kiểm soát domain shift, mất cân bằng lớp và chất lượng landmark.",
        bold=True,
    )
    add_bullets(
        doc,
        [
            "VSL có chuyển động lớn hơn ASL: total_motion cao hơn khoảng 2.07 lần, motion_variance cao hơn 2.17 lần, mean_velocity cao hơn 1.51 lần và direction_change_count cao hơn 1.65 lần.",
            "VSL dài hơn rõ rệt: trung bình 92.5 frame so với 65.6 frame của ASL; max VSL lên tới 465 frame.",
            "VSL rất thưa mẫu theo nhãn: median chỉ 1 mẫu/nhãn, 2,763 nhãn chỉ có 1 mẫu, gần như toàn bộ nhãn dưới 5 mẫu. Đây là rủi ro lớn nhất cho supervised classification.",
            "ASL ổn định hơn về số mẫu/nhãn nhưng vẫn thưa ở chuẩn deep learning: median 6 mẫu/nhãn và 1,896/2,000 nhãn dưới 10 mẫu.",
            "Một số nhãn/cặp nhãn có DTW rất thấp hoặc bằng 0, nghĩa là quỹ đạo tay giống nhau hoặc dữ liệu tay bị mất, dễ gây nhầm lớp nếu chỉ dựa vào trajectory.",
        ],
    )

    counts = video.groupby("dataset").agg(
        videos=("file", "count"),
        labels=("label", "nunique"),
        mean_frames=("num_frames", "mean"),
        median_frames=("num_frames", "median"),
        p95_frames=("num_frames", lambda s: s.quantile(0.95)),
        max_frames=("num_frames", "max"),
    ).reset_index()
    add_heading(doc, "2. Quy mô và độ dài chuỗi", 1)
    rows = [[r.dataset, fmt(r.videos, 0), fmt(r.labels, 0), fmt(r.mean_frames), fmt(r.median_frames), fmt(r.p95_frames), fmt(r.max_frames, 0)] for r in counts.itertuples()]
    add_table(doc, ["Dataset", "Video", "Nhãn", "Frame TB", "Median", "P95", "Max"], rows, widths=[0.8, 0.9, 0.8, 0.9, 0.8, 0.8, 0.8], font_size=9)
    add_para(
        doc,
        "Ý nghĩa khi train: model cần xử lý chuỗi dài không đồng nhất. Nếu padding/truncation quá ngắn, VSL bị mất pha chuyển động cuối; nếu padding quá dài, ASL bị nhiều padding, làm batch kém hiệu quả. Nên dùng mask temporal, bucketing theo độ dài, hoặc resampling động thay vì ép cứng một độ dài cho mọi mẫu.",
    )
    add_figure(doc, "boxplot_num_frames.png", "Boxplot số frame: VSL dài hơn và có đuôi phân phối lớn hơn ASL.")
    add_figure(doc, "histogram_num_frames.png", "Histogram số frame: phân phối độ dài là cơ sở chọn max_len, bucketing và sampling.")

    add_heading(doc, "3. So sánh đặc trưng chuyển động ASL - VSL", 1)
    metrics = [
        ("num_frames", "Số frame", "VSL dài hơn, cần mask/bucketing."),
        ("hand_frame_missing_ratio", "Tỷ lệ frame thiếu tay", "ASL cao hơn, nhưng cả hai đều cần xử lý missing landmark."),
        ("total_motion", "Tổng chuyển động", "VSL lớn hơn 2.07x, biểu hiện domain motion mạnh."),
        ("mean_velocity", "Vận tốc TB", "VSL nhanh hơn 1.51x."),
        ("motion_variance", "Phương sai chuyển động", "VSL dao động/không đều hơn 2.17x."),
        ("direction_change_count", "Số đổi hướng", "VSL nhiều đổi hướng hơn 1.65x."),
        ("straightness_ratio", "Độ thẳng quỹ đạo", "ASL cao hơn; VSL nhiều đường cong/lặp hơn."),
        ("left_right_hand_dist_mean", "Khoảng cách 2 tay", "VSL cao hơn, cần giữ quan hệ hai tay."),
        ("hand_face_dist_mean", "Tay - mặt", "VSL cao hơn; đặc trưng mặt/thân có thể hữu ích."),
        ("hand_body_dist_mean", "Tay - thân", "VSL cao hơn nhẹ."),
    ]
    mean_by_ds = video.groupby("dataset").mean(numeric_only=True)
    rows = []
    for key, label, note in metrics:
        asl = mean_by_ds.loc["ASL", key]
        vsl = mean_by_ds.loc["VSL", key]
        diff = vsl - asl
        ratio = vsl / asl if asl not in [0, math.nan] and not pd.isna(asl) else math.nan
        rows.append([label, fmt(asl), fmt(vsl), fmt(diff), f"{ratio:.2f}x" if not pd.isna(ratio) else "", note])
    add_table(doc, ["Metric", "ASL", "VSL", "VSL-ASL", "Tỷ lệ", "Diễn giải"], rows, widths=[1.25, 0.65, 0.65, 0.75, 0.6, 2.6], font_size=8)
    add_para(
        doc,
        "Ảnh hưởng trực tiếp: nếu train mô hình chung ASL+VSL, model có thể học tín hiệu dataset thay vì học ngữ nghĩa ký hiệu. Ví dụ, chỉ cần nhìn độ dài/biên độ chuyển động cũng có thể đoán domain. Khi mục tiêu là nhận dạng trong từng ngôn ngữ, nên train/tune riêng; khi mục tiêu là mô hình đa ngôn ngữ, cần domain-aware split, domain embedding hoặc adversarial/domain-balanced training.",
    )
    for fn, cap in [
        ("boxplot_total_motion.png", "Total motion: VSL có biên độ chuyển động cao hơn."),
        ("boxplot_mean_velocity.png", "Mean velocity: VSL nhanh hơn, ảnh hưởng tới temporal augmentation."),
        ("boxplot_motion_variance.png", "Motion variance: VSL có nhiều biến thiên, khó ổn định hơn."),
        ("boxplot_straightness_ratio.png", "Straightness ratio: ASL có quỹ đạo thẳng hơn, VSL cong/lặp nhiều hơn."),
        ("boxplot_left_right_hand_dist_mean.png", "Khoảng cách hai tay: cần mô hình hóa tương tác hai tay thay vì chỉ dùng một tay."),
        ("boxplot_trajectory_length.png", "Trajectory length gần tương đương total_motion và nhấn mạnh khác biệt biên độ."),
        ("boxplot_mean_acceleration.png", "Acceleration: VSL có độ biến thiên chuyển động cao hơn."),
    ]:
        add_figure(doc, fn, cap)

    add_heading(doc, "4. Chất lượng landmark và missing data", 1)
    quality = video.groupby("dataset").agg(
        hand_missing=("hand_frame_missing_ratio", "mean"),
        left_missing=("left_frame_missing_ratio", "mean"),
        right_missing=("right_frame_missing_ratio", "mean"),
        pose_missing=("pose_frame_missing_ratio", "mean"),
        face_missing=("face_frame_missing_ratio", "mean"),
        usable_motion=("usable_motion", "mean"),
    ).reset_index()
    rows = [[r.dataset, pct(r.hand_missing), pct(r.left_missing), pct(r.right_missing), pct(r.pose_missing), pct(r.face_missing), pct(r.usable_motion)] for r in quality.itertuples()]
    add_table(doc, ["Dataset", "Thiếu tay", "Thiếu tay trái", "Thiếu tay phải", "Thiếu pose", "Thiếu face", "Usable"], rows, widths=[0.75, 0.8, 0.95, 0.95, 0.8, 0.8, 0.75], font_size=9)
    add_para(
        doc,
        "Tỷ lệ thiếu tay là tín hiệu rủi ro cho mọi mô hình dùng skeleton/keypoint. ASL thiếu tay trung bình 39.6%, VSL 32.3%. Pose/face gần như đầy đủ, do đó tay mới là bottleneck chính. Các nhãn VSL 'bất động', 'cười nhếch mép', 'gật gù', 'hắt hơi', 'lắc đầu' có hand_missing_mean = 100%, làm DTW giữa chúng bằng 0; nếu bài toán cần nhận dạng các nhãn này, phải dùng face/head/body hoặc dữ liệu RGB, không thể chỉ dựa vào hand trajectory.",
    )
    missing_rows = []
    for ds in ["ASL", "VSL"]:
        for r in by_label[by_label.dataset.eq(ds)].sort_values("hand_missing_mean", ascending=False).head(8).itertuples():
            missing_rows.append([ds, r.label, fmt(r.num_samples, 0), pct(r.hand_missing_mean), fmt(r.mean_total_motion)])
    add_table(doc, ["Dataset", "Nhãn thiếu tay cao", "Mẫu", "Thiếu tay", "Total motion TB"], missing_rows, widths=[0.75, 2.1, 0.55, 0.75, 0.95], font_size=8)

    add_heading(doc, "5. Mất cân bằng lớp và rủi ro supervised learning", 1)
    label_dist = by_label.groupby("dataset").agg(
        labels=("label", "count"),
        videos=("num_samples", "sum"),
        median_samples=("num_samples", "median"),
        mean_samples=("num_samples", "mean"),
        min_samples=("num_samples", "min"),
        max_samples=("num_samples", "max"),
        labels_lt2=("num_samples", lambda s: (s < 2).sum()),
        labels_lt5=("num_samples", lambda s: (s < 5).sum()),
        labels_lt10=("num_samples", lambda s: (s < 10).sum()),
    ).reset_index()
    rows = [[r.dataset, fmt(r.labels, 0), fmt(r.videos, 0), fmt(r.median_samples), fmt(r.mean_samples), fmt(r.min_samples, 0), fmt(r.max_samples, 0), fmt(r.labels_lt2, 0), fmt(r.labels_lt5, 0), fmt(r.labels_lt10, 0)] for r in label_dist.itertuples()]
    add_table(doc, ["Dataset", "Nhãn", "Video", "Median", "Mean", "Min", "Max", "<2", "<5", "<10"], rows, widths=[0.7, 0.6, 0.75, 0.65, 0.65, 0.45, 0.45, 0.45, 0.45, 0.45], font_size=8)
    add_para(
        doc,
        "Đây là điểm ảnh hưởng mạnh nhất đến accuracy. Với VSL, đa số nhãn chỉ có một mẫu nên không thể học classifier đa lớp ổn định theo cách supervised thông thường. Train/test split ngẫu nhiên theo video cũng dễ tạo nhãn không xuất hiện ở train hoặc chỉ xuất hiện ở test. Với ASL, số mẫu khá hơn nhưng vẫn thấp cho 2,000 lớp; cần class-balanced sampling và đánh giá top-k/macro-F1 thay vì chỉ accuracy.",
    )

    add_heading(doc, "6. Nhãn chuyển động phức tạp", 1)
    complex_rows = []
    for ds in ["ASL", "VSL"]:
        for r in by_label[by_label.dataset.eq(ds)].sort_values("motion_complexity_score", ascending=False).head(10).itertuples():
            complex_rows.append([ds, r.label, fmt(r.num_samples, 0), fmt(r.mean_total_motion), fmt(r.mean_velocity), fmt(r.motion_variance_mean), fmt(r.direction_change_mean), pct(r.hand_missing_mean), fmt(r.motion_complexity_score)])
    add_table(doc, ["DS", "Nhãn", "Mẫu", "Motion", "Velocity", "Variance", "Đổi hướng", "Thiếu tay", "Score"], complex_rows, widths=[0.45, 1.7, 0.45, 0.65, 0.65, 0.65, 0.65, 0.65, 0.55], font_size=7)
    add_para(
        doc,
        "Các nhãn phức tạp nên được dùng làm nhóm kiểm thử riêng. Tuy nhiên với VSL, nhiều nhãn top complexity chỉ có 1 mẫu, nên score chủ yếu chỉ ra outlier hoặc mẫu khó, chưa đủ để kết luận nhãn đó thật sự khó trên toàn bộ phân phối.",
    )
    add_figure(doc, "top_complex_motion_labels_ASL.png", "Top nhãn ASL có chuyển động phức tạp.")
    add_figure(doc, "top_complex_motion_labels_VSL.png", "Top nhãn VSL có chuyển động phức tạp; cần đọc cùng số mẫu/nhãn.")
    add_figure(doc, "scatter_total_motion_vs_motion_variance.png", "Scatter total_motion vs motion_variance: nhận diện outlier chuyển động mạnh.")
    add_figure(doc, "scatter_trajectory_length_vs_straightness_ratio.png", "Scatter trajectory_length vs straightness_ratio: phân biệt đường thẳng, cong và lặp.")

    add_heading(doc, "7. DTW nội lớp và cặp lớp dễ nhầm", 1)
    intra_rows = []
    for ds, g in intra.groupby("dataset"):
        intra_rows.append([ds, fmt(len(g), 0), fmt(g.intra_class_dtw_mean.mean()), fmt(g.intra_class_dtw_mean.median()), fmt(g.intra_class_dtw_mean.quantile(0.75)), fmt(g.intra_class_dtw_mean.quantile(0.90)), fmt(g.intra_class_dtw_mean.max())])
    add_table(doc, ["Dataset", "Nhãn có DTW", "Mean", "Median", "P75", "P90", "Max"], intra_rows, widths=[0.8, 1.0, 0.75, 0.75, 0.65, 0.65, 0.65], font_size=9)
    inter_rows = []
    for ds, g in inter.groupby("dataset"):
        inter_rows.append([ds, fmt(len(g), 0), fmt(g.dtw_distance.min()), fmt(g.dtw_distance.quantile(0.10)), fmt(g.dtw_distance.median()), fmt((g.dtw_distance <= 1e-9).sum(), 0)])
    add_table(doc, ["Dataset", "Cặp ứng viên", "Min", "P10", "Median", "Cặp DTW=0"], inter_rows, widths=[0.8, 1.2, 0.65, 0.65, 0.75, 0.8], font_size=9)
    add_para(
        doc,
        "DTW nội lớp cao nghĩa là cùng một nhãn được thực hiện khác nhau giữa các mẫu/người ký hiệu; model cần học invariance theo tốc độ, biên độ và pha chuyển động. DTW liên lớp thấp nghĩa là hai nhãn có quỹ đạo giống nhau; nếu chỉ dùng tay, classifier dễ nhầm và cần thêm đặc trưng vị trí tương đối với mặt/thân, hướng tay, face/head cue hoặc loss phân biệt cặp hard negative.",
    )
    pair_rows = []
    for ds in ["ASL", "VSL"]:
        for r in inter[inter.dataset.eq(ds)].head(10).itertuples():
            pair_rows.append([ds, r.label_a, r.label_b, fmt(r.dtw_distance, 4)])
    add_table(doc, ["DS", "Nhãn A", "Nhãn B", "DTW"], pair_rows, widths=[0.45, 2.0, 2.0, 0.6], font_size=8)
    add_figure(doc, "histogram_intra_class_dtw.png", "Histogram DTW nội lớp: độ biến thiên trong cùng nhãn.")
    add_figure(doc, "top_inter_class_dtw_pairs_ASL.png", "Các cặp ASL có DTW thấp, dễ nhầm khi chỉ dựa vào quỹ đạo.")
    add_figure(doc, "top_inter_class_dtw_pairs_VSL.png", "Các cặp VSL có DTW thấp; các cặp bằng 0 thường liên quan đến thiếu landmark tay.")

    add_heading(doc, "8. Diễn giải biểu đồ quỹ đạo", 1)
    add_para(
        doc,
        "Các ảnh trajectory_examples cho thấy khác biệt hình học của quỹ đạo tay sau khi resample. Chúng hữu ích để kiểm tra trực quan xem normalization có giữ được hình chuyển động hay không. Nếu trajectory bị co về một điểm hoặc đường rất ngắn trong khi nhãn cần chuyển động tay, nên kiểm tra lại missing-hand handling và normalization_scale.",
    )
    add_figure(doc, "trajectory_examples_ASL.png", "Ví dụ quỹ đạo ASL sau resampling.")
    add_figure(doc, "trajectory_examples_VSL.png", "Ví dụ quỹ đạo VSL sau resampling.")
    for fn, cap in [
        ("histogram_total_motion.png", "Histogram total_motion: phân phối biên độ chuyển động."),
        ("histogram_mean_velocity.png", "Histogram mean_velocity: phân phối tốc độ ký hiệu."),
        ("histogram_motion_variance.png", "Histogram motion_variance: phân phối độ bất ổn chuyển động."),
    ]:
        add_figure(doc, fn, cap)

    add_heading(doc, "9. Khuyến nghị train mô hình", 1)
    add_heading(doc, "9.1. Tách bài toán trước khi tối ưu mô hình", 2)
    add_bullets(
        doc,
        [
            "Nếu mục tiêu là nhận dạng ASL và VSL riêng: train/tune hai mô hình hoặc ít nhất hai head phân loại riêng, vì phân phối chuyển động và số mẫu/nhãn khác nhau mạnh.",
            "Nếu mục tiêu là mô hình chung: thêm domain_id ASL/VSL, dùng balanced sampler theo domain, và báo cáo metric riêng từng domain. Không chỉ báo cáo accuracy gộp.",
            "Với VSL hiện tại, không nên kỳ vọng classifier closed-set hàng nghìn lớp đạt tốt nếu giữ tất cả nhãn 1 mẫu. Cần lọc nhãn tối thiểu 3-5 mẫu hoặc chuyển sang few-shot/metric learning.",
        ],
    )

    add_heading(doc, "9.2. Preprocessing nên áp dụng", 2)
    add_bullets(
        doc,
        [
            "Dùng temporal mask thực sự cho padding; không để model học padding như tín hiệu lớp/domain.",
            "Chuẩn hóa tọa độ theo thân người hoặc shoulder/torso scale, nhưng giữ thêm đặc trưng khoảng cách tay-mặt, tay-thân và hai tay vì chúng có giá trị phân biệt.",
            "Xử lý missing hand bằng mask channel riêng, forward-fill ngắn hạn, interpolation có kiểm soát; không thay NaN/thiếu tay bằng 0 mà không báo mask cho model.",
            "Dùng resampling/augmentation thời gian: speed perturbation, random crop temporal, time masking. Với VSL cần giữ chuỗi dài hoặc dùng model có attention/mask.",
            "Lọc hoặc gắn cờ mẫu hand_missing quá cao, đặc biệt các nhãn VSL thiếu tay 100%. Các nhãn phi-tay cần nhánh face/head/body hoặc dữ liệu RGB.",
        ],
    )

    add_heading(doc, "9.3. Chiến lược mô hình", 2)
    add_bullets(
        doc,
        [
            "Baseline mạnh: Transformer/TCN trên skeleton sequence với attention mask, input gồm hand+pose+face/mouth features và missing masks.",
            "Với lớp ít mẫu: dùng metric learning/prototypical network, supervised contrastive loss, hoặc pretrain self-supervised trên toàn bộ sequence rồi fine-tune lớp đủ mẫu.",
            "Dùng class-balanced loss hoặc focal loss; sampler cân bằng theo class và domain. Với ASL, oversample nhãn ít mẫu; với VSL, cân nhắc loại nhãn 1 mẫu khỏi supervised split chính.",
            "Thêm hard-negative mining từ `dtw_inter_class_candidate_pairs.csv`, ví dụ cute-sweet, sugar-voice, bao tay-găng tay, sổ tay-sổ điểm danh.",
            "Đánh giá bằng macro-F1, balanced accuracy, top-5 accuracy, confusion matrix theo nhãn khó và metric riêng cho nhãn high_missing/high_DTW.",
        ],
    )

    add_heading(doc, "9.4. Split dữ liệu và kiểm thử", 2)
    add_bullets(
        doc,
        [
            "Split theo nhãn chỉ khi làm few-shot/open-set; split theo video trong cùng nhãn chỉ hợp lệ nếu mỗi nhãn có đủ mẫu train/val/test.",
            "Không dùng random split đơn giản cho VSL toàn bộ vì nhiều nhãn chỉ có 1 mẫu. Nên tạo subset VSL đủ mẫu để supervised classification và giữ phần còn lại cho retrieval/few-shot.",
            "Tạo test set 'hard labels': nhãn motion_complexity cao, sequence_length_variance cao, intra_class_dtw cao, và các cặp inter_class_dtw thấp.",
            "Báo cáo thêm robustness theo missing-hand bucket: thấp, trung bình, cao. Nếu model tụt mạnh ở bucket missing cao, cần cải thiện extractor hoặc input modality.",
        ],
    )

    add_heading(doc, "10. Kế hoạch hành động đề xuất", 1)
    action_rows = [
        ("1", "Làm sạch dữ liệu", "Gắn cờ/loại mẫu hand_missing quá cao; kiểm tra nhãn VSL có hand_missing=100%; giữ face/head/body nếu nhãn không dùng tay.", "Giảm nhiễu label và feature."),
        ("2", "Tạo subset train hợp lệ", "ASL: giữ toàn bộ nhưng dùng balanced sampler. VSL: tạo subset nhãn >=3 hoặc >=5 mẫu cho supervised; nhãn 1 mẫu chuyển sang few-shot/retrieval.", "Tránh train classifier trên lớp không đủ dữ liệu."),
        ("3", "Baseline skeleton", "Transformer/TCN có temporal mask, missing mask, feature gồm tay + pose + face/mouth + khoảng cách tương đối.", "Có baseline mạnh và giải thích được."),
        ("4", "Hard negative training", "Dùng cặp DTW thấp để tăng contrastive/hard-negative loss.", "Giảm nhầm các nhãn quỹ đạo giống nhau."),
        ("5", "Đánh giá tách domain", "Báo cáo ASL riêng, VSL riêng, macro-F1, top-k, confusion nhóm nhãn khó.", "Không che lấp lỗi bằng accuracy gộp."),
    ]
    add_table(doc, ["Bước", "Việc cần làm", "Cách làm", "Tác dụng"], action_rows, widths=[0.45, 1.25, 3.1, 1.7], font_size=8)

    add_heading(doc, "11. Phụ lục nguồn đã tổng hợp", 1)
    source_rows = [
        ("motion_features_per_video.csv", f"{len(video):,} dòng", "Đặc trưng từng video: motion, velocity, missing, distances."),
        ("motion_summary_by_dataset.csv / wide", f"{len(wide):,} metric", "So sánh trung bình ASL/VSL."),
        ("motion_features_by_label.csv", f"{len(by_label):,} dòng", "Thống kê theo nhãn và complexity score."),
        ("sequence_length_by_label.csv", f"{len(seq):,} dòng", "Độ dài chuỗi theo nhãn."),
        ("top_complex_motion_labels.csv", f"{len(top_complex):,} dòng", "Top nhãn chuyển động phức tạp."),
        ("dtw_intra_class_summary.csv", f"{len(intra):,} dòng", "Biến thiên nội lớp bằng DTW."),
        ("dtw_inter_class_candidate_pairs.csv", f"{len(inter):,} dòng", "Cặp nhãn dễ nhầm theo DTW thấp."),
        ("debug/npz_structure_inspection.csv", f"{len(debug):,} dòng", "Kiểm tra cấu trúc file NPZ đầu vào."),
        ("figures/*.png", f"{len(list(FIG.glob('*.png')))} ảnh", "Boxplot, histogram, scatter, trajectory, top labels/pairs."),
    ]
    add_table(doc, ["Nguồn", "Quy mô", "Đã dùng để"], source_rows, widths=[2.3, 1.0, 3.2], font_size=8)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
